"""
DOCX Exporter - Generates Microsoft Word documents.
"""
from pathlib import Path
from typing import Optional, Dict, Any, List
import re

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from ..parser.notebook_parser import ParsedNotebook, NotebookCell
from ..graph_handler.graph_storage import StoredGraph


class DOCXExporter:
    """Exports formatted notebooks to Word documents."""
    
    def __init__(self):
        self.document: Optional[Document] = None
    
    def export(
        self,
        notebook: ParsedNotebook,
        output_path: Path,
        graphs: Optional[List[StoredGraph]] = None,
        options: Optional[Dict[str, Any]] = None,
    ) -> Path:
        """Export notebook to DOCX format."""
        options = options or {}
        
        # Create new document
        self.document = Document()
        
        # Set up styles
        self._setup_styles()
        
        # Add title
        self._add_title(notebook.title)
        
        # Add table of contents placeholder
        if options.get("toc", True):
            self._add_toc_placeholder()
        
        # Process cells
        for cell in notebook.cells:
            self._add_cell(cell)
        
        # Add figures section
        if graphs:
            self._add_figures_section(graphs)
        
        # Save document
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.document.save(str(output_path))
        
        return output_path
    
    def _setup_styles(self):
        """Set up document styles."""
        styles = self.document.styles
        
        # Code style
        if "Code" not in [s.name for s in styles]:
            code_style = styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
            code_style.font.name = "Consolas"
            code_style.font.size = Pt(9)
            code_style.paragraph_format.space_before = Pt(6)
            code_style.paragraph_format.space_after = Pt(6)
        
        # Output style
        if "Output" not in [s.name for s in styles]:
            output_style = styles.add_style("Output", WD_STYLE_TYPE.PARAGRAPH)
            output_style.font.name = "Consolas"
            output_style.font.size = Pt(8)
            output_style.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    
    def _add_title(self, title: str):
        """Add document title."""
        heading = self.document.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_toc_placeholder(self):
        """Add table of contents placeholder."""
        self.document.add_paragraph("Table of Contents")
        self.document.add_paragraph(
            "[Update this field to generate table of contents]",
            style="Intense Quote"
        )
        self.document.add_page_break()
    
    def _add_cell(self, cell: NotebookCell):
        """Add a notebook cell to the document."""
        if cell.cell_type == "markdown":
            self._add_markdown(cell.source)
        elif cell.cell_type == "code":
            self._add_code(cell.source)
            self._add_outputs(cell.outputs)
    
    def _add_markdown(self, content: str):
        """Add markdown content (converted to Word formatting)."""
        lines = content.split("\n")
        
        for line in lines:
            stripped = line.strip()
            
            # Handle headings
            if stripped.startswith("### "):
                self.document.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                self.document.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                self.document.add_heading(stripped[2:], level=1)
            # Handle lists
            elif stripped.startswith("- ") or stripped.startswith("* "):
                self.document.add_paragraph(stripped[2:], style="List Bullet")
            elif re.match(r"^\d+\.\s", stripped):
                text = re.sub(r"^\d+\.\s", "", stripped)
                self.document.add_paragraph(text, style="List Number")
            # Handle bold/italic (simplified)
            elif stripped:
                para = self.document.add_paragraph()
                self._add_formatted_text(para, stripped)
    
    def _add_formatted_text(self, paragraph, text: str):
        """Add text with basic markdown formatting."""
        # Simple bold/italic handling
        parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
        
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith("`") and part.endswith("`"):
                run = paragraph.add_run(part[1:-1])
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            else:
                paragraph.add_run(part)
    
    def _add_code(self, code: str):
        """Add code block to document."""
        # Add a subtle border effect with a table
        table = self.document.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        
        cell = table.rows[0].cells[0]
        
        # Add code line by line
        lines = code.split("\n")
        for i, line in enumerate(lines):
            para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            para.style = "Code" if "Code" in [s.name for s in self.document.styles] else None
            run = para.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
    
    def _add_outputs(self, outputs: list):
        """Add cell outputs to document."""
        for output in outputs:
            if output.is_image:
                # Skip - images go to figures section
                continue
            
            if output.text:
                # Truncate long outputs
                text = output.text
                if len(text) > 1000:
                    text = text[:1000] + "\n... (output truncated)"
                
                para = self.document.add_paragraph(style="Output")
                run = para.add_run(text)
                run.font.name = "Consolas"
                run.font.size = Pt(8)
    
    def _add_figures_section(self, graphs: List[StoredGraph]):
        """Add figures section at the end."""
        self.document.add_page_break()
        self.document.add_heading("Figures", level=1)
        
        for graph in graphs:
            file_path = Path(graph.file_path)
            if not file_path.exists():
                continue
            
            # Add figure
            self.document.add_picture(
                str(file_path),
                width=Inches(5.5)
            )
            
            # Add caption
            caption = self.document.add_paragraph(graph.caption)
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.style = "Caption"
            
            # Add some space
            self.document.add_paragraph()
    
    def export_from_html(
        self,
        html_content: str,
        output_path: Path,
        options: Optional[Dict[str, Any]] = None,
    ) -> Path:
        """Export HTML content to DOCX (limited support)."""
        # This is a simplified converter
        # For full HTML support, consider using pandoc or mammoth
        
        self.document = Document()
        
        # Extract text content (very basic)
        text = re.sub(r'<[^>]+>', '\n', html_content)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        for para in text.split('\n\n'):
            para = para.strip()
            if para:
                self.document.add_paragraph(para)
        
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.document.save(str(output_path))
        
        return output_path
